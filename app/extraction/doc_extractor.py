from __future__ import annotations

import io
from zipfile import BadZipFile

from docx import Document


class DocExtractor:
    """Извлекает текст из офисных документов форматов DOCX/DOC."""

    def extract(self, raw_file: bytes, extension: str) -> tuple[str, int, bool]:
        """Извлекает текст из DOCX; для DOC сообщает о необходимости предварительной конвертации."""
        normalized = extension.lower()
        if normalized == ".docx":
            doc = Document(io.BytesIO(raw_file))
            text = "\n".join(par.text.strip() for par in doc.paragraphs if par.text).strip()
            return text, 1, False

        if normalized == ".doc":
            raise RuntimeError(
                "DOC binary format is not supported directly. Convert DOC -> DOCX via LibreOffice/Tika before retry."
            )

        raise RuntimeError(f"Unsupported office extension: {extension}")

    @staticmethod
    def is_docx_bytes(raw_file: bytes) -> bool:
        """Проверяет, можно ли интерпретировать переданные байты как валидный DOCX."""
        try:
            Document(io.BytesIO(raw_file))
            return True
        except BadZipFile:
            return False
